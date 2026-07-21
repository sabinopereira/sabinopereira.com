#!/usr/bin/env python3
from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "books/borrow-delay-repeat/source/borrow-delay-repeat-source.docx"
OUT = ROOT / "books/borrow-delay-repeat/directors-cut"
DOCX = OUT / "borrow-delay-repeat-directors-cut.docx"

INK = RGBColor(25, 24, 22)
RED = RGBColor(158, 34, 38)
MUTED = RGBColor(105, 99, 91)
PAPER = "F6F0E4"

COMMENTARY = {
    "Chapter 1 — Ask Nicely": (
        "The first trick is not the request. It is the casting. The lender is offered a flattering role: dependable, generous, different from everyone who has ever disappointed you. Once a person accepts the role, saying no can feel like bad acting.",
        "On the other side of the screen, kindness is doing paperwork. The person reading the message is checking their balance, rearranging a bill, and deciding whether your discomfort matters more than theirs. Their yes may be generous. It is not weightless.",
        "Every loan begins as money. The dishonest ones quickly become theatre."
    ),
    "Chapter 2 — Promise Tomorrow": (
        "Tomorrow works because it is close enough to picture and far enough to escape. The narrator is not selling a date; he is selling relief from having to discuss the date.",
        "The lender does not hear a poetic concept. They hear a calendar entry. They may plan rent, transport, food, or their own promise around yours. A false deadline does not pause the problem. It quietly transfers it.",
        "A promise can calm a conversation while making the debt more expensive."
    ),
    "Chapter 3 — Strategic Disappearance": (
        "Selective silence is funny until the absent person remains visibly active everywhere else. Then the disappearance becomes a performance with push notifications.",
        "No reply creates work. The lender must decide whether to ask again, soften the wording, pretend not to care, or risk looking desperate for requesting what already belongs to them.",
        "Silence does not erase a debt. It makes the other person carry both sides of the conversation."
    ),
    "Chapter 4 — Image Control": (
        "The narrator treats appearances as evidence. If he looks relaxed, perhaps the account is relaxed too. But a restaurant photo posted beside an unanswered message is not neutral scenery. It is an accidental financial statement.",
        "The lender rarely resents joy. They resent being asked to finance the illusion that nothing is wrong. Public ease beside private avoidance turns uncertainty into disrespect.",
        "The image may stay polished. The credit underneath it does not."
    ),
    "Chapter 5 — The Follow-up Defence": (
        "Here the comedy darkens. A normal question is reframed as aggression so the borrower can become the injured party. The debt remains, but now the lender is also expected to apologise for remembering it.",
        "Following up already feels awkward. Being punished for it teaches the lender a clean lesson: next time, protect the relationship by refusing the loan.",
        "When accountability feels like an attack, trust starts planning its exit."
    ),
    "Chapter 6 — The Reset": (
        "Charm is used as a solvent. Enough warmth, enough nostalgia, enough confidence—and perhaps the previous scene will dissolve. It does not. It only becomes the subtext of the new one.",
        "A friendly return can feel like pressure to choose between peace and truth. The lender may smile, not because the account is settled, but because they have stopped expecting settlement.",
        "A reset without repair is just repetition wearing clean clothes."
    ),
    "Final Chapter — Legacy": (
        "The final account is not measured in currency. It appears in delayed replies, invitations that never arrive, and emergencies nobody quite believes. The narrator wanted to borrow money without consequence; he ended up lending doubt to every future version of himself.",
        "Most people do not announce that trust is gone. They simply redesign their life so it is no longer required. The calls get shorter. The boundaries get clearer. The wallet stays closed.",
        "You can owe money once. If you repeat the pattern, people begin to think you owe them the truth."
    ),
}

def font(run, name="Georgia", size=11, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def shade(paragraph, fill=PAPER):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)

def clean_source():
    raw = [p.text.strip() for p in Document(SOURCE).paragraphs if p.text.strip()]
    fixed = []
    for text in raw:
        if " Part 1 — The Basics" in text:
            left, _ = text.split(" Part 1 — The Basics", 1)
            if left.strip(): fixed.append(left.strip())
            fixed.append("Part I — The Borrow")
        elif text.startswith("Chapter 3 — Strategic Disappearance"):
            fixed += ["Chapter 3 — Strategic Disappearance", text.replace("Chapter 3 — Strategic Disappearance", "").strip()]
        elif text.startswith("Chapter 5 — The Follow-up Defense"):
            fixed += ["Chapter 5 — The Follow-up Defence", text.replace("Chapter 5 — The Follow-up Defense", "").strip()]
        elif text not in {raw[0], raw[1]}:
            fixed.append(re.sub(r"\s+", " ", text))
    replacements = {
        "Civilized.": "Civilised.",
        "defense": "defence",
        "Defense": "Defence",
        "It’s not avoidance.": "This is not avoidance.",
        "Incase": "In case",
    }
    cleaned = []
    for text in fixed:
        for old, new in replacements.items():
            text = text.replace(old, new)
        text = re.sub(r"\s+([,.;:!?])", r"\1", text)
        text = re.sub(r"\.{2,}", "…", text)
        cleaned.append(text.strip())
    return [x for x in cleaned if x]

def source_sections():
    markers = {"Disclaimer", "Intro", "Part I — The Borrow"}
    result = []
    current = None
    for text in clean_source():
        is_heading = text in markers or text.startswith("Chapter ") or text.startswith("Final Chapter")
        if is_heading:
            current = {"title": text, "lines": []}
            result.append(current)
        elif current:
            current["lines"].append(text)
    return result

def prose_paragraphs(lines):
    """Recompose line-broken source into readable prose while retaining rare comic beats."""
    deliberate_beats = {
        "One word.",
        "Instant peace.",
        "You don’t panic.",
        "They do.",
        "That’s how you know it’s working.",
        "Borrow. Delay. Repeat.",
    }
    paragraphs, buffer = [], []

    def flush():
        if buffer:
            text = " ".join(buffer)
            text = re.sub(r"\s+", " ", text).strip()
            paragraphs.append(("prose", text))
            buffer.clear()

    for line in lines:
        if line in deliberate_beats:
            flush()
            paragraphs.append(("beat", line))
            continue
        buffer.append(line)
        words = sum(len(item.split()) for item in buffer)
        strong_end = line.endswith((".", "?", "!", "…", ".”", "?’", "!’"))
        if words >= 65 and strong_end:
            flush()
    flush()
    return paragraphs

def add_recomposed_body(doc, lines):
    for kind, text in prose_paragraphs(lines):
        p = doc.add_paragraph(text)
        if kind == "beat":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(9)
            if p.runs:
                font(p.runs[0], "Georgia", 10.6, italic=True, color=RED)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)

def configure(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(6), Inches(9)
    sec.top_margin, sec.bottom_margin = Inches(.8), Inches(.78)
    sec.left_margin, sec.right_margin = Inches(.78), Inches(.72)
    sec.header_distance, sec.footer_distance = Inches(.35), Inches(.35)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size, normal.font.color.rgb = "Georgia", Pt(10.6), INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    for style_name, size, color, before, after in [("Title", 30, INK, 0, 10), ("Subtitle", 13, MUTED, 0, 8), ("Heading 1", 20, INK, 18, 9), ("Heading 2", 12, RED, 14, 6)]:
        s = doc.styles[style_name]
        s.font.name, s.font.size, s.font.color.rgb = "Georgia", Pt(size), color
        s.font.bold = style_name != "Subtitle"
        s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
        s.paragraph_format.keep_with_next = True
    if "Receipt" not in [s.name for s in doc.styles]:
        s = doc.styles.add_style("Receipt", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name, s.font.size, s.font.color.rgb = "Courier New", Pt(9.5), INK
        s.paragraph_format.left_indent = Inches(.18)
        s.paragraph_format.right_indent = Inches(.18)
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after = Pt(10)
    header = sec.header.paragraphs[0]
    header.text = "BORROW. DELAY. REPEAT.  /  DIRECTOR’S CUT"
    font(header.runs[0], "Courier New", 7.5, bold=True, color=MUTED)
    add_page_number(sec.footer.paragraphs[0])

def add_cover(doc):
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BORROW.\nDELAY.\nREPEAT.")
    font(r, "Arial", 31, bold=True)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph("DIRECTOR’S CUT", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("A Very Unserious Guide to Not Giving Money Back")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.runs[0], "Georgia", 12, italic=True, color=MUTED)
    doc.add_paragraph()
    p = doc.add_paragraph("SABINO PEREIRA")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.runs[0], "Arial", 10, bold=True, color=RED)
    doc.add_page_break()

def add_callout(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.18)
    p.paragraph_format.right_indent = Inches(.18)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    shade(p)
    r = p.add_run(title.upper() + "\n")
    font(r, "Arial", 8.5, bold=True, color=RED)
    r = p.add_run(body)
    font(r, "Georgia", 10, italic=False)

def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    add_cover(doc)
    doc.add_heading("A Note Before the Excuses", 1)
    doc.add_paragraph("This book is satire. The narrator is confident, articulate, and frequently wrong. His methods are presented so the pattern can be recognised—not copied. If someone trusted you with money, the least expensive response is usually the honest one.")
    add_callout(doc, "Terms and conditions", "No friendship is improved by strategic silence. No screenshot counts as a payment plan. ‘Tomorrow’ is a date, not a personality.")
    doc.add_page_break()

    for section in source_sections():
        text = section["title"]
        if text == "Disclaimer":
            doc.add_heading("Disclaimer", 1)
        elif text == "Intro":
            doc.add_heading("Introduction — The Flexible Conscience", 1)
        elif text == "Part I — The Borrow":
            doc.add_page_break(); doc.add_heading(text, 1)
            doc.add_paragraph("The request is small. The role assigned to the other person is not.", style="Subtitle")
        elif text.startswith("Chapter ") or text.startswith("Final Chapter"):
            title = text
            if title == "Chapter 3 — Strategic Disappearance":
                doc.add_page_break(); doc.add_heading("Part II — The Delay", 1)
            if title == "Chapter 6 — The Reset":
                doc.add_page_break(); doc.add_heading("Part III — The Repeat", 1)
            doc.add_page_break(); doc.add_heading(title, 1)

        add_recomposed_body(doc, section["lines"])

        if text in COMMENTARY:
            commentary, other_side, receipt = COMMENTARY[text]
            add_callout(doc, "Director’s commentary", commentary)
            doc.add_heading("From the Other Side", 2)
            doc.add_paragraph(other_side)
            p = doc.add_paragraph(style="Receipt")
            shade(p, "EEE5D6")
            p.add_run("THE RECEIPT\n" + receipt)

    doc.add_page_break()
    doc.add_heading("How to Break the Loop", 1)
    doc.add_paragraph("The joke ends when the conversation becomes specific. Name the amount. Name what happened without decorating it. Offer a date you can meet, or a sequence of smaller payments you can actually sustain. If the date changes, say so before the other person has to chase you.")
    doc.add_paragraph("An honest message can be short: ‘I owe you ____. I cannot pay all of it today. I can send ____ on ____ and ____ on ____. If that changes, I will tell you before the date.’ Then do the uncinematic part: keep the promise.")
    doc.add_heading("For the Person Being Asked", 2)
    doc.add_paragraph("You are allowed to say no. You are allowed to ask for the amount and date in writing. You are allowed to lend only what you could survive never seeing again. A boundary is not a verdict on someone’s character; it is a decision about your own capacity.")
    p = doc.add_paragraph(style="Receipt"); shade(p, "EEE5D6")
    p.add_run("FINAL BALANCE\nMoney can be repaid. Trust returns through evidence, one kept promise at a time.")
    doc.save(DOCX)
    print(DOCX)

if __name__ == "__main__":
    build()
