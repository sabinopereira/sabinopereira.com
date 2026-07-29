#!/usr/bin/env python3
from pathlib import Path
import html
import re
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate, Flowable, Frame, KeepTogether, PageBreak,
    PageTemplate, Paragraph, Spacer
)

ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript" / "a-journey-back-to-silence-final-manuscript.md"
COVER = ROOT / "assets" / "cover" / "a-journey-back-to-silence-book-cover.png"
OUTPUT_PDF = ROOT / "output" / "pdf" / "a-journey-back-to-silence-reading-edition.pdf"
OUTPUT_DOCX = ROOT / "output" / "a-journey-back-to-silence-final-manuscript.docx"

PAGE_W, PAGE_H = 6 * inch, 9 * inch
PAPER = HexColor("#F3EDE3")
NAVY = HexColor("#102A3A")
OCEAN = HexColor("#486B7B")
MIST = HexColor("#9AA9AE")
INK = HexColor("#262B2D")
GOLD = HexColor("#A58662")

FONT_REG = "/System/Library/Fonts/Supplemental/Georgia.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Georgia Bold.ttf"
FONT_ITALIC = "/System/Library/Fonts/Supplemental/Georgia Italic.ttf"
FONT_BOLD_ITALIC = "/System/Library/Fonts/Supplemental/Georgia Bold Italic.ttf"

for name, path in {
    "Georgia": FONT_REG,
    "Georgia-Bold": FONT_BOLD,
    "Georgia-Italic": FONT_ITALIC,
    "Georgia-BoldItalic": FONT_BOLD_ITALIC,
}.items():
    pdfmetrics.registerFont(TTFont(name, path))
pdfmetrics.registerFontFamily(
    "Georgia", normal="Georgia", bold="Georgia-Bold",
    italic="Georgia-Italic", boldItalic="Georgia-BoldItalic"
)


def clean_inline(value):
    value = html.escape(value.strip(), quote=False)
    value = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", value)
    value = re.sub(r"\*(.+?)\*", r"<i>\1</i>", value)
    return value


def blocks(text):
    return [item.strip() for item in re.split(r"\n\s*\n", text.strip()) if item.strip()]


def split_top_sections(text):
    matches = list(re.finditer(r"(?m)^# (.+)$", text))
    result = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        result.append((match.group(1).strip(), text[match.end():end].strip()))
    return result


class FullPageImage(Flowable):
    def __init__(self, path):
        super().__init__()
        self.path = str(path)

    def wrap(self, avail_width, avail_height):
        return 0, 0

    def drawOn(self, canvas, x, y, _sW=0):
        canvas.drawImage(ImageReader(self.path), 0, 0, PAGE_W, PAGE_H,
                         preserveAspectRatio=False, mask="auto")


class DropCapOpening(Flowable):
    """Print-edition opening paragraph with a three-line classical initial."""
    def __init__(self, text):
        super().__init__()
        self.letter = text[0]
        self.remainder = text[1:].lstrip()
        self.cap_width = 27
        self.gap = 3
        self.paragraph = None
        self.height = 0

    def wrap(self, avail_width, avail_height):
        opening_style = ParagraphStyle(
            "DropCapText", parent=BODY, alignment=TA_JUSTIFY,
            spaceAfter=0, leading=14.8,
        )
        self.paragraph = Paragraph(clean_inline(self.remainder), opening_style)
        _, paragraph_height = self.paragraph.wrap(
            avail_width - self.cap_width - self.gap, avail_height
        )
        self.width = avail_width
        self.height = max(paragraph_height, 38)
        return avail_width, self.height

    def draw(self):
        self.canv.saveState()
        self.canv.setFillColor(NAVY)
        self.canv.setFont("Georgia", 35)
        self.canv.drawString(0, self.height - 31, self.letter)
        self.paragraph.drawOn(
            self.canv, self.cap_width + self.gap,
            self.height - self.paragraph.height,
        )
        self.canv.restoreState()


class SilenceBookDoc(BaseDocTemplate):
    def __init__(self, filename):
        super().__init__(
            filename,
            pagesize=(PAGE_W, PAGE_H),
            leftMargin=0.78 * inch,
            rightMargin=0.70 * inch,
            topMargin=0.72 * inch,
            bottomMargin=0.72 * inch,
            title="A Journey Back to Silence",
            author="Sabino Pereira",
            subject="Eight Passages from Noise to Stillness",
        )
        frame = Frame(
            self.leftMargin, self.bottomMargin,
            PAGE_W - self.leftMargin - self.rightMargin,
            PAGE_H - self.topMargin - self.bottomMargin,
            id="body", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
        )
        self.addPageTemplates([PageTemplate(id="book", frames=frame, onPage=draw_page)])


def draw_page(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(PAPER)
    canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    if doc.page > 5:
        canvas.setStrokeColor(HexColor("#D8D0C5"))
        canvas.setLineWidth(0.35)
        canvas.line(0.78 * inch, 0.53 * inch, PAGE_W - 0.70 * inch, 0.53 * inch)
        canvas.setFillColor(OCEAN)
        canvas.setFont("Georgia", 7.7)
        canvas.drawCentredString(PAGE_W / 2, 0.31 * inch, str(doc.page - 4))
    canvas.restoreState()


styles = getSampleStyleSheet()
BODY = ParagraphStyle(
    "BookBody", fontName="Georgia", fontSize=10.15, leading=14.8,
    textColor=INK, alignment=TA_JUSTIFY, spaceAfter=8.2,
    allowWidows=0, allowOrphans=0,
)
BODY_FIRST = ParagraphStyle("BookBodyFirst", parent=BODY)
MAJOR = ParagraphStyle(
    "Major", fontName="Georgia", fontSize=25, leading=30,
    textColor=NAVY, alignment=TA_CENTER, spaceAfter=18,
)
KICKER = ParagraphStyle(
    "Kicker", fontName="Georgia", fontSize=8.2, leading=11,
    textColor=GOLD, alignment=TA_CENTER, spaceAfter=13,
)
SUBTITLE = ParagraphStyle(
    "Subtitle", fontName="Georgia-Italic", fontSize=11.5, leading=17,
    textColor=OCEAN, alignment=TA_CENTER, spaceAfter=18,
)
SECTION = ParagraphStyle(
    "Section", fontName="Georgia", fontSize=12.5, leading=16,
    textColor=NAVY, alignment=TA_LEFT, spaceBefore=15, spaceAfter=8,
)
QUESTION = ParagraphStyle(
    "Question", fontName="Georgia-Italic", fontSize=11.2, leading=17,
    textColor=NAVY, alignment=TA_LEFT, leftIndent=12, rightIndent=12,
    spaceBefore=8, spaceAfter=12,
)
PAUSE = ParagraphStyle(
    "Pause", fontName="Georgia-Italic", fontSize=12, leading=18,
    textColor=OCEAN, alignment=TA_CENTER, leftIndent=22, rightIndent=22,
    spaceBefore=14, spaceAfter=10,
)
CONTENTS = ParagraphStyle(
    "Contents", fontName="Georgia", fontSize=9.4, leading=15.2,
    textColor=INK, alignment=TA_LEFT, leftIndent=18,
)
SMALL = ParagraphStyle(
    "Small", fontName="Georgia", fontSize=8.4, leading=12.5,
    textColor=OCEAN, alignment=TA_LEFT,
)


def add_pdf_prose(story, section_body, use_drop_cap=False):
    current_subhead = None
    drop_cap_pending = False
    for block in blocks(section_body):
        if block == "---":
            continue
        if block.startswith("### "):
            current_subhead = block[4:].strip()
            if use_drop_cap and current_subhead == "The Arrival":
                drop_cap_pending = True
            story.append(Paragraph(clean_inline(current_subhead), SECTION))
        elif block.startswith("## "):
            continue
        elif block.startswith("> "):
            quote = " ".join(line.lstrip("> ") for line in block.splitlines())
            story.append(Spacer(1, 0.18 * inch))
            story.append(Paragraph(clean_inline(quote), PAUSE))
        elif block.startswith("**") and block.endswith("**"):
            story.append(Paragraph(clean_inline(block), SUBTITLE))
        else:
            style = QUESTION if current_subhead == "A Quiet Question" else BODY
            plain_block = block.replace("  \n", "<br/>")
            if drop_cap_pending and current_subhead == "The Arrival":
                story.append(DropCapOpening(plain_block))
                story.append(Spacer(1, 8.2))
                drop_cap_pending = False
            else:
                story.append(Paragraph(clean_inline(plain_block), style))


def build_pdf():
    text = MANUSCRIPT.read_text(encoding="utf-8")
    sections = split_top_sections(text)
    story = [FullPageImage(COVER), PageBreak()]

    story.extend([
        Spacer(1, 1.28 * inch),
        Paragraph("A JOURNEY", MAJOR),
        Paragraph("BACK TO SILENCE", MAJOR),
        Paragraph("EIGHT PASSAGES FROM NOISE TO STILLNESS", KICKER),
        Spacer(1, 0.42 * inch),
        Paragraph("Written by Sabino Pereira", SUBTITLE),
        Paragraph("Companion music by Reira Bin", SUBTITLE),
        PageBreak(),
        Spacer(1, 1.35 * inch),
        Paragraph("Copyright © 2026 Sabino Pereira", SMALL),
        Spacer(1, 0.15 * inch),
        Paragraph("All rights reserved. No part of this publication may be reproduced, stored or transmitted in any form without prior written permission, except for brief quotations used in reviews.", SMALL),
        Spacer(1, 0.15 * inch),
        Paragraph("First reading edition, 2026", SMALL),
        Spacer(1, 0.15 * inch),
        Paragraph("Written by Sabino Pereira (SP). Companion music composed by Reira Bin (RB).", SMALL),
        PageBreak(),
    ])

    for name, body in sections:
        if name == "A JOURNEY BACK TO SILENCE":
            continue
        if name == "Dedication":
            story.extend([Spacer(1, 1.65 * inch), Paragraph("DEDICATION", KICKER)])
            for block in blocks(body):
                if block != "---":
                    story.append(Paragraph(clean_inline(block), SUBTITLE))
            story.append(PageBreak())
            continue
        if name == "A Note to the Reader":
            story.extend([Spacer(1, 1.1 * inch), Paragraph(name, MAJOR)])
            add_pdf_prose(story, body)
            story.append(PageBreak())
            continue
        if name == "Contents":
            story.extend([Spacer(1, 0.55 * inch), Paragraph("Contents", MAJOR)])
            for line in body.replace("---", "").splitlines():
                line = line.strip()
                if line:
                    story.append(Paragraph(clean_inline(line), CONTENTS))
            story.append(PageBreak())
            continue
        if name in {"Part I", "Part II", "Part III"}:
            title = re.search(r"(?m)^## (.+)$", body).group(1)
            italic = re.search(r"(?m)^\*(.+)\*$", body)
            story.extend([
                Spacer(1, 1.75 * inch),
                Paragraph(name.upper(), KICKER),
                Paragraph(clean_inline(title), MAJOR),
                Spacer(1, 0.15 * inch),
            ])
            if italic:
                story.append(Paragraph(clean_inline(f"*{italic.group(1)}*"), SUBTITLE))
            story.append(PageBreak())
            continue
        if name.startswith("Passage "):
            title_match = re.search(r"(?m)^## (.+)$", body)
            title = title_match.group(1) if title_match else name
            remaining = re.sub(r"(?m)^## .+\n?", "", body, count=1).strip()
            story.extend([
                Spacer(1, 0.48 * inch),
                Paragraph(name.upper(), KICKER),
                Paragraph(clean_inline(title), MAJOR),
                Spacer(1, 0.08 * inch),
            ])
            add_pdf_prose(story, remaining, use_drop_cap=True)
            story.append(PageBreak())
            continue
        if name == "Final Page":
            story.extend([Spacer(1, 2.0 * inch)])
            add_pdf_prose(story, body)
            continue
        story.extend([Spacer(1, 0.58 * inch), Paragraph(clean_inline(name), MAJOR)])
        add_pdf_prose(story, body)
        story.append(PageBreak())

    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    doc = SilenceBookDoc(str(OUTPUT_PDF))
    doc.build(story)


def set_run_font(run, name="Georgia", size=11, bold=False, italic=False, color="262B2D"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, end])
    set_run_font(run, size=8, color="486B7B")


def style_docx(doc):
    # narrative_proposal preset, with named book overrides: 6x9 trim and Georgia typography.
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("262B2D")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, "102A3A"),
        ("Heading 2", 13, 12, 6, "486B7B"),
        ("Heading 3", 12, 8, 4, "486B7B"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    add_page_number(section.footer.paragraphs[0])


def add_docx_paragraph(doc, text, style=None, align=None, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, italic=italic)
    return p


def build_docx():
    text = MANUSCRIPT.read_text(encoding="utf-8")
    sections = split_top_sections(text)
    doc = Document()
    style_docx(doc)
    props = doc.core_properties
    props.title = "A Journey Back to Silence"
    props.subject = "Eight Passages from Noise to Stillness"
    props.author = "Sabino Pereira"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    r = p.add_run("A JOURNEY BACK TO SILENCE")
    set_run_font(r, size=26, color="102A3A")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EIGHT PASSAGES FROM NOISE TO STILLNESS")
    set_run_font(r, size=10, color="A58662")
    add_docx_paragraph(doc, "Written by Sabino Pereira", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_docx_paragraph(doc, "Companion music by Reira Bin", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_page_break()

    add_docx_paragraph(doc, "Copyright © 2026 Sabino Pereira")
    add_docx_paragraph(doc, "All rights reserved. First reading edition, 2026.")
    doc.add_page_break()

    for name, body in sections:
        if name == "A JOURNEY BACK TO SILENCE":
            continue
        if name == "Contents":
            doc.add_heading(name, level=1)
            for line in body.replace("---", "").splitlines():
                line = line.strip()
                if line:
                    p = add_docx_paragraph(doc, line.replace("  ", ""))
                    p.paragraph_format.space_after = Pt(3)
            doc.add_page_break()
            continue
        if name in {"Part I", "Part II", "Part III"}:
            title_match = re.search(r"(?m)^## (.+)$", body)
            italic_match = re.search(r"(?m)^\*(.+)\*$", body)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(150)
            r = p.add_run(name.upper())
            set_run_font(r, size=10, color="A58662")
            if title_match:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(title_match.group(1))
                set_run_font(r, size=22, color="102A3A")
            if italic_match:
                add_docx_paragraph(doc, italic_match.group(1), align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
            doc.add_page_break()
            continue
        if name.startswith("Passage "):
            title_match = re.search(r"(?m)^## (.+)$", body)
            doc.add_heading(name, level=1)
            if title_match:
                doc.add_heading(title_match.group(1), level=2)
                body = re.sub(r"(?m)^## .+\n?", "", body, count=1).strip()
        else:
            doc.add_heading(name, level=1)
        for block in blocks(body):
            if block == "---":
                continue
            if block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=2)
            elif block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=3)
            elif block.startswith("> "):
                quote = " ".join(line.lstrip("> ") for line in block.splitlines())
                p = add_docx_paragraph(doc, quote, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.right_indent = Inches(0.35)
            elif block.startswith("*") and block.endswith("*"):
                add_docx_paragraph(doc, block.strip("*"), align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
            else:
                plain = block.replace("**", "").replace("  \n", "\n")
                add_docx_paragraph(doc, plain)
        doc.add_page_break()

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main():
    if not COVER.exists():
        raise FileNotFoundError(f"Missing cover: {COVER}")
    build_docx()
    build_pdf()
    print(OUTPUT_DOCX)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    main()
