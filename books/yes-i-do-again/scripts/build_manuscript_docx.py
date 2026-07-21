from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "yes-i-do-again.md"
OUTPUT_DIR = ROOT / "output"
OUTPUT = OUTPUT_DIR / "yes-i-do-again-manuscript.docx"

# Base preset: narrative_proposal. Named override: trade_fiction_manuscript.
# The override replaces proposal typography with a restrained book-manuscript system.
PAGE = {"width": 8.5, "height": 11.0, "margin": 1.0, "header": 0.492, "footer": 0.492}
FONT = "Garamond"
BODY_SIZE = 11
BODY_LINE = 1.10
INK = RGBColor(0x18, 0x18, 0x18)
MUTED = RGBColor(0x66, 0x66, 0x66)


def set_font(run, name=FONT, size=BODY_SIZE, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE["width"])
    section.page_height = Inches(PAGE["height"])
    section.top_margin = Inches(PAGE["margin"])
    section.bottom_margin = Inches(PAGE["margin"])
    section.left_margin = Inches(PAGE["margin"])
    section.right_margin = Inches(PAGE["margin"])
    section.header_distance = Inches(PAGE["header"])
    section.footer_distance = Inches(PAGE["footer"])

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Inches(0.28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = BODY_LINE
    pf.widow_control = True

    for style_name, size, before, after in (
        ("Heading 1", 18, 0, 18),
        ("Heading 2", 15, 0, 16),
        ("Heading 3", 12, 12, 8),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = INK
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if style_name == "Heading 1":
            style.paragraph_format.page_break_before = True

    if "Front Matter" not in [s.name for s in doc.styles]:
        front = doc.styles.add_style("Front Matter", WD_STYLE_TYPE.PARAGRAPH)
    else:
        front = doc.styles["Front Matter"]
    front.font.name = FONT
    front._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    front._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    front.font.size = Pt(11)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.space_after = Pt(8)


def add_inline_markdown(paragraph, text):
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith("*") and part.endswith("*")
        value = part[1:-1] if italic else part
        run = paragraph.add_run(value)
        set_font(run, italic=italic)


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph(style="Normal")
    # Dialogue is still prose, but a new speaker receives a modest first-line indent.
    add_inline_markdown(paragraph, text)
    return paragraph


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Title page: editorial_cover pattern, simplified for fiction.
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("YES, I DO... AGAIN.")
    set_font(r, size=28)
    p = doc.add_paragraph(style="Front Matter")
    r = p.add_run("A Novel")
    set_font(r, size=12, italic=True, color=MUTED)
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Front Matter")
    r = p.add_run("SABINO PEREIRA")
    set_font(r, size=14)

    doc.add_page_break()
    p = doc.add_paragraph(style="Front Matter")
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Copyright © 2026 Sabino Pereira")
    set_font(r, size=10)
    p = doc.add_paragraph(style="Front Matter")
    r = p.add_run("All rights reserved.")
    set_font(r, size=10)
    p = doc.add_paragraph(style="Front Matter")
    r = p.add_run("This is a work of fiction. Names, characters, places, and incidents are products of the author’s imagination or are used fictitiously.")
    set_font(r, size=9.5, italic=True, color=MUTED)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    first_title_skipped = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# ") and not first_title_skipped:
            first_title_skipped = True
            continue
        if stripped.startswith("## "):
            heading_text = stripped[3:]
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(54)
            r = p.add_run(heading_text.upper() if heading_text.startswith("Chapter") else heading_text)
            set_font(r, size=18)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(stripped[4:])
            set_font(r, size=15)
            continue
        add_body_paragraph(doc, stripped)

    # Front matter has no page number; body sections share footer in this compact source doc.
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = "Yes, I Do... Again."
    doc.core_properties.author = "Sabino Pereira"
    doc.core_properties.subject = "English-language novel manuscript"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
