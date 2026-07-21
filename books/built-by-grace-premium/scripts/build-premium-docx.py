from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "build" / "built-by-grace-premium-manuscript.md"
OUTPUT = ROOT / "dist" / "Built by Grace - Premium Edition.docx"
COVER = ROOT / "assets" / "built-by-grace-premium-cover-v2.png"

INK = RGBColor(43, 31, 24)
GOLD = RGBColor(166, 118, 46)
MUTED = RGBColor(104, 84, 66)
BODY_FONT = "Garamond"
DISPLAY_FONT = "Georgia"


def font(run, name=BODY_FONT, size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_rich_text(paragraph, text, size=11, color=INK, italic=False, bold=False):
    pieces = re.split(r"(\*[^*]+\*)", text)
    for piece in pieces:
        if not piece:
            continue
        is_em = piece.startswith("*") and piece.endswith("*")
        run = paragraph.add_run(piece[1:-1] if is_em else piece)
        font(run, size=size, color=color, italic=italic or is_em, bold=bold)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, GOLD),
        ("Heading 2", 13, 12, 6, GOLD),
        ("Heading 3", 12, 8, 4, MUTED),
    ):
        style = doc.styles[style_name]
        style.font.name = DISPLAY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DISPLAY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DISPLAY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    section = doc.sections[0]
    section.page_width = Inches(6.25)
    section.page_height = Inches(10)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.add_run().add_picture(str(COVER), width=Inches(6.25), height=Inches(10))

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(6.25)
    section.page_height = Inches(10)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("BUILT BY GRACE"), DISPLAY_FONT, 26, GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    font(p.add_run("A Journey of Prayer, Healing, Love,\nand God in the Center"), DISPLAY_FONT, 13, MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    font(p.add_run("—  ✦  —"), DISPLAY_FONT, 12, GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("SABINO PEREIRA"), DISPLAY_FONT, 11, INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(92)
    font(p.add_run("PREMIUM DIGITAL EDITION"), DISPLAY_FONT, 8.5, MUTED)
    return section


def add_contents(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("CONTENTS"), DISPLAY_FONT, 10.5, MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run("The Journey"), DISPLAY_FONT, 23, GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run("—  ✦  —"), DISPLAY_FONT, 10, GOLD)

    entries = [
        ("Introduction", "Before Love, There Was a Prayer"),
        ("Chapter 1", "The Prayer Before Love"),
        ("Chapter 2", "A Heart Prepared"),
        ("Chapter 3", "It’s Safe to Love Again"),
        ("Chapter 4", "Love Shouldn’t Feel Like Fear"),
        ("Chapter 5", "Teach Us Love"),
        ("Chapter 6", "You Stayed"),
        ("Chapter 7", "God in the Center"),
        ("Chapter 8", "Healing Is Coming — Grace Is Still Fighting"),
        ("Chapter 9", "Mercy Over Fear — Let Mercy Rise"),
        ("Chapter 10", "God Sent Me You"),
        ("Chapter 11", "No Silent Walls"),
        ("Chapter 12", "Goodness and Mercy — Mercy Found Me"),
        ("Chapter 13", "Family Prayer — Built by Grace"),
        ("Chapter 14", "The Home Prayer"),
        ("Chapter 15", "Built by Grace"),
        ("Bonus Reflection", "Before We Say I Do"),
        ("Closing", "God in the Center: Final Blessing"),
    ]
    for label, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
        font(p.add_run(f"{label.upper()}  "), DISPLAY_FONT, 8.5, GOLD, bold=True)
        font(p.add_run(title), BODY_FONT, 10, INK)

    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    setup_styles(doc)

    section = add_cover(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(header.add_run("BUILT BY GRACE  ·  PREMIUM EDITION"), DISPLAY_FONT, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("—  "), DISPLAY_FONT, 9, MUTED)
    add_field(footer, "PAGE")
    font(footer.add_run("  —"), DISPLAY_FONT, 9, MUTED)

    raw_blocks = [b.strip() for b in re.split(r"\n{2,}", SOURCE.read_text(encoding="utf-8")) if b.strip()]
    blocks = []
    for block in raw_blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) > 1 and all(line.startswith("#") for line in lines):
            blocks.extend(lines)
        else:
            blocks.append(block)
    prayer_mode = False
    closing_thought_mode = False
    chapter_heading_pending = False
    chapter_body_pending = False
    special_section = ""
    contents_added = False
    skipped_title = 0

    for block in blocks:
        if block == "\\pagebreak":
            doc.add_page_break()
            prayer_mode = False
            closing_thought_mode = False
            continue
        if block.startswith("# Built by Grace") or block.startswith("## A Journey") or block == "by Sabino Pereira":
            skipped_title += 1
            continue
        if block.startswith("# "):
            special_section = ""
            prayer_mode = block.startswith("# Closing")
            closing_thought_mode = False
            chapter_heading_pending = True
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(76)
            p.paragraph_format.space_after = Pt(8)
            add_rich_text(p, block[2:].strip().upper(), size=10.5, color=MUTED, bold=True)
            continue
        if block.startswith("## "):
            if not chapter_heading_pending:
                special_section = block[3:].strip()
            p = doc.add_paragraph(style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if chapter_heading_pending:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(14)
                add_rich_text(p, block[3:].strip(), size=22, color=GOLD, bold=True)
                ornament = doc.add_paragraph()
                ornament.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ornament.paragraph_format.space_after = Pt(22)
                font(ornament.add_run("—  ✦  —"), DISPLAY_FONT, 10, GOLD)
                chapter_heading_pending = False
                chapter_body_pending = True
            else:
                add_rich_text(p, block[3:].strip(), size=16, color=GOLD, bold=True)
            continue
        if block.startswith("### "):
            if block.startswith("### Introduction"):
                special_section = ""
                if not contents_added:
                    add_contents(doc)
                    contents_added = True
                chapter_heading_pending = True
                p = doc.add_paragraph(style="Heading 1")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(76)
                p.paragraph_format.space_after = Pt(8)
                add_rich_text(p, "INTRODUCTION", size=10.5, color=MUTED, bold=True)
                continue
            prayer_mode = block.startswith("### Prayer") or block.startswith("### Closing Thought")
            closing_thought_mode = block.startswith("### Closing Thought")
            p = doc.add_paragraph(style="Heading 3")
            if block.startswith("### Prayer"):
                p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, block[4:].strip(), size=12, color=MUTED, bold=True)
            continue

        if chapter_body_pending:
            doc.add_page_break()
            chapter_body_pending = False

        p = doc.add_paragraph()
        if prayer_mode:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.2
            if closing_thought_mode:
                p.paragraph_format.keep_with_next = True
            lines = block.splitlines()
            for index, line in enumerate(lines):
                add_rich_text(
                    p,
                    line.strip().replace("  ", ""),
                    size=12 if closing_thought_mode else 11,
                    color=MUTED if closing_thought_mode else INK,
                    italic=closing_thought_mode,
                )
                if index < len(lines) - 1:
                    p.add_run().add_break(WD_BREAK.LINE)
        else:
            if special_section == "Copyright":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_after = Pt(10)
            elif special_section in {"Dedication", "Premium Edition"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_after = Pt(14)
            elif special_section == "Author's Note":
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Inches(0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.22)
            add_rich_text(p, block.replace("\n", " "), size=11)

    doc.core_properties.title = "Built by Grace — Premium Edition"
    doc.core_properties.subject = "A Journey of Prayer, Healing, Love, and God in the Center"
    doc.core_properties.author = "Sabino Pereira"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
