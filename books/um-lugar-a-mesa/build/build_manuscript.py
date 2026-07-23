from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Um Lugar à Mesa - Manuscrito Reconstruído.docx"
SOURCES = [
    ROOT / "source/front-matter.md",
    ROOT / "source/capitulo-01-completo.md",
    ROOT / "source/capitulos-02-08.md",
    ROOT / "source/capitulos-09-epilogo.md",
    ROOT / "source/epilogo-completo.md",
]

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Garamond"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.18
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.first_line_indent = Inches(0.22)

for name, size in [("Title", 30), ("Heading 1", 22), ("Heading 2", 14)]:
    st = styles[name]
    st.font.name = "Garamond"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor(35, 28, 24)
    st.font.bold = name != "Title"
    st.paragraph_format.space_before = Pt(18)
    st.paragraph_format.space_after = Pt(12)

if "Epigraph" not in styles:
    st = styles.add_style("Epigraph", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Garamond"
    st.font.size = Pt(12)
    st.font.italic = True
    st.paragraph_format.left_indent = Inches(1.0)
    st.paragraph_format.right_indent = Inches(1.0)
    st.paragraph_format.space_after = Pt(10)

# Cover image and title page
cover = ROOT / "assets/um-lugar-a-mesa-cover-premium-1600x2560.jpg"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(cover), height=Inches(8.2))
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(150)
r = p.add_run("UM LUGAR À MESA")
r.font.name = "Garamond"
r.font.size = Pt(30)
r.bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("SABINO PEREIRA")
r.font.name = "Garamond"
r.font.size = Pt(16)
doc.add_page_break()

def clean_inline(text):
    return text.replace("**", "").replace("*", "")

def prose_lines(text):
    """Reflow sentence-per-line source into natural novel paragraphs."""
    output = []
    buffer = []
    buffer_chars = 0

    def flush():
        nonlocal buffer, buffer_chars
        if buffer:
            output.append(" ".join(buffer))
            buffer = []
            buffer_chars = 0

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        control = line.startswith("#") or line == "---" or re.match(r"^\d+\. ", line)
        dialogue = line.startswith("“") or line.startswith('"')
        emphasis = (line.startswith("**") and line.endswith("**")) or (
            line.startswith("*") and line.endswith("*")
        )
        if control or dialogue or emphasis:
            flush()
            output.append(line)
            continue
        buffer.append(line)
        buffer_chars += len(line)
        if len(buffer) >= 4 or buffer_chars >= 420 or line.endswith(":"):
            flush()
    flush()
    return output

chapter_section_started = False
skipped_front_separator = False
for source in SOURCES:
    source_text = source.read_text(encoding="utf-8")
    if source.name == "capitulos-02-08.md":
        full_chapter_02 = (ROOT / "source/capitulo-02-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 2\n.*?(?=# CAPÍTULO 3)",
            full_chapter_02 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_03 = (ROOT / "source/capitulo-03-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 3\n.*?(?=# CAPÍTULO 4)",
            full_chapter_03 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_04 = (ROOT / "source/capitulo-04-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 4\n.*?(?=# CAPÍTULO 5)",
            full_chapter_04 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_05 = (ROOT / "source/capitulo-05-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 5\n.*?(?=# CAPÍTULO 6)",
            full_chapter_05 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_06 = (ROOT / "source/capitulo-06-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 6\n.*?(?=# CAPÍTULO 7)",
            full_chapter_06 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_07 = (ROOT / "source/capitulo-07-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 7\n.*?(?=# CAPÍTULO 8)",
            full_chapter_07 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_08 = (ROOT / "source/capitulo-08-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 8\n.*\Z",
            full_chapter_08 + "\n",
            source_text,
            flags=re.S,
        )
    if source.name == "capitulos-09-epilogo.md":
        full_chapter_09 = (ROOT / "source/capitulo-09-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 9\n.*?(?=# CAPÍTULO 10)",
            full_chapter_09 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_10 = (ROOT / "source/capitulo-10-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 10\n.*?(?=# CAPÍTULO 11)",
            full_chapter_10 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_11 = (ROOT / "source/capitulo-11-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 11\n.*?(?=# CAPÍTULO 12)",
            full_chapter_11 + "\n\n",
            source_text,
            flags=re.S,
        )
        full_chapter_12 = (ROOT / "source/capitulo-12-completo.md").read_text(encoding="utf-8").strip()
        source_text = re.sub(
            r"# CAPÍTULO 12\n.*?(?=# EPÍLOGO)",
            full_chapter_12 + "\n\n",
            source_text,
            flags=re.S,
        )
    prepared_lines = source_text.splitlines() if source.name == "front-matter.md" else prose_lines(source_text)
    for raw in prepared_lines:
        line = raw.strip()
        if not line:
            continue
        if source.name == "capitulos-09-epilogo.md" and line == "# EPÍLOGO":
            break
        if line == "---":
            if source.name == "epilogo-completo.md":
                p = doc.add_paragraph("◆")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                continue
            if source.name == "front-matter.md" and not skipped_front_separator:
                skipped_front_separator = True
                continue
            doc.add_page_break()
            continue
        if line.startswith("# "):
            title = clean_inline(line[2:])
            if source.name == "front-matter.md" and title == "UM LUGAR À MESA":
                continue
            if not chapter_section_started:
                chapter_sec = doc.add_section(WD_SECTION.NEW_PAGE)
                chapter_sec.top_margin = Inches(0.85)
                chapter_sec.bottom_margin = Inches(0.85)
                chapter_sec.left_margin = Inches(0.9)
                chapter_sec.right_margin = Inches(0.9)
                chapter_sec.footer.is_linked_to_previous = False
                sect_pr = chapter_sec._sectPr
                pg_num = OxmlElement("w:pgNumType")
                pg_num.set(qn("w:start"), "1")
                sect_pr.append(pg_num)
                fp = chapter_sec.footer.paragraphs[0]
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fld = OxmlElement("w:fldSimple")
                fld.set(qn("w:instr"), "PAGE")
                fp._p.append(fld)
                chapter_section_started = True
            else:
                doc.add_page_break()
            p = doc.add_paragraph(title, style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            continue
        if line.startswith("## "):
            text = clean_inline(line[3:])
            if source.name == "front-matter.md" and text == "Sabino Pereira":
                continue
            p = doc.add_paragraph(text, style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(clean_inline(line), style="Normal")
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.25)
            continue
        if line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph(clean_inline(line), style="Epigraph")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        p = doc.add_paragraph(style="Normal")
        p.add_run(clean_inline(line))

doc.core_properties.title = "Um Lugar à Mesa"
doc.core_properties.author = "Sabino Pereira"
doc.core_properties.subject = "Manuscrito reconstruído para revisão editorial"
doc.save(OUT)
print(OUT)
